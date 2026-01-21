import wmg_feedback_gen.core as core
from docxtpl import DocxTemplate
import openpyxl
import jinja2
from docx import Document
from io import BytesIO
from jinja2 import meta
import logging
import os
import docx.table
import docx.enum.text

def highlight_cell(cell: docx.table._Cell, color=docx.enum.text.WD_COLOR_INDEX.YELLOW):
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.highlight_color = color

def default_hightlight(row_data: dict, filename: str):
    logger = logging.getLogger(__name__)

    logger.debug(f"Post-processing file: {filename}")

    docx = Document(filename)
    for table in docx.tables:
        logger.debug(f"Processing table with {len(table.rows)} rows.")
        for row in table.rows:
            logger.debug(f"Processing row with {len(row.cells)} cells.")
            if len(row.cells) not in (9,8): # not the correct table
                continue

            logger.debug(f"Row cells: {[cell.text for cell in row.cells]}")

            categories = ["OUTSTANDING", "DISTINCTION", "GOOD", "PASS", "MARGINAL", "FAIL"]

            if len(row.cells) == 9: # has the KSB column
                logger.debug("Detected KSB column in the table.")
                lookup = dict(zip(categories, row.cells[3:9]))
                comments = row.cells[2]
            else: # no KSB column
                lookup = dict(zip(categories, row.cells[2:8]))
                comments = row.cells[1]

            category = comments.text.strip().split()[-1]
            logger.debug(f"Identified category: {category}")
            
            # Highlight the category in the document
            try:
                highlight_cell(lookup[category])
                logger.debug(f"Highlighting {category} in {filename}")
            except KeyError:
                continue

    # Save the modified document to the output file
    docx.save(filename)

def generate_doc(row_data: dict, 
                 template: DocxTemplate, 
                 output_filename: str,
                 jinja_env=None,
                 post_processing=default_hightlight):
    template.reset_replacements()
    template.render(row_data, jinja_env=jinja_env)

    # Save to an in-memory file
    template.save(output_filename)

    # Call post_processing after saving the document if it's a function
    if callable(post_processing) and post_processing.__code__.co_argcount == 2:
        post_processing(row_data, output_filename)
    else:
        logging.debug("Post-processing function absent or does not match expected signature.")


def generate( 
    xlsx_filename: str,
    template_filename: str,
    worksheet: str = "marks",
    output_filename: str = "feedback/feedback_{{STUDENTID}}.docx",
    validators: dict = core.default_validators,
    jinga_env=None,
    post_processing=default_hightlight,
    expected_vars=None):
    """
    Find the columns in the given worksheet that match the expected variable names.

    Args:
        xlsx_filename: The Excel worksbook to open.
        template_filename: The filename of the Word document to use as template.
        worksheet: The name of the worksheet to process in the Excel workbook.
        output_filename: The filename pattern for the output documents. 
        validators: A dictionary of validators for the variables.
        jinga_env: An optional Jinja2 environment to use for rendering.
        post_processing: A function to call after generating each document.
        expected_vars: A set of expected variable names to look for in the worksheet.

    Returns:
        

    Details:
        Validators are functions that are run against the relevant column name to test
        if this row should generate an output document.

        Post-processing is a function that can be used to modify generated documents
        after they have been created. It is called with the row data and the filename
        of the generated document.

        Expected variables are the set of column names that are expected to be present in the worksheet.
        For the most part these will be automatically extracted from the Word template, output filename, 
        and validators. 
        However, if using the post_processing function, it may be necessary to specify additional variables
    """

    tpl = DocxTemplate(template_filename)

    if jinga_env is None:
        # Create a new Jinja2 environment if not provided
        jinja_env = jinja2.Environment()
        jinja_env.filters['mark_category'] = core.mark_category

    if validators is None:
        validators = {}

    # Extract Jinja variables from the output_filename string
    output_filename_ast = jinja_env.parse(output_filename)
    output_filename_vars = meta.find_undeclared_variables(output_filename_ast)

    # Extract undeclared template variables from the template
    template_vars = tpl.get_undeclared_template_variables(jinja_env=jinja_env)

    variables = template_vars.union(output_filename_vars).union(validators.keys())
    if expected_vars is not None:
        for var in expected_vars: variables.add(var)

    workbook = openpyxl.load_workbook(xlsx_filename, data_only=True)

    # create feedback directory if it does not exist
    os.makedirs(os.path.dirname(output_filename), exist_ok=True)

    for row_data in core.process_to_dicts(workbook[worksheet], variables, validators):
        #logging.debug(f"Processing row data: {row_data}")
        generate_doc(row_data,
                     tpl,
                     core.gen_filename(output_filename, row_data),
                     jinja_env=jinja_env,
                     post_processing=post_processing)

